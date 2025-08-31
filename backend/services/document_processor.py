import os
import uuid
from typing import List, Dict, Any
from pathlib import Path
import asyncio
import logging
from datetime import datetime

# Document processing imports
try:
    import pypdf
    from pypdf import PdfReader
except ImportError:
    try:
        import PyPDF2
        from PyPDF2 import PdfReader
    except ImportError:
        print("Warning: No PDF library available")
        PdfReader = None

from docx import Document as DocxDocument
import spacy

# AI models
from transformers import pipeline
from sentence_transformers import SentenceTransformer

from models.schemas import Document, DocumentChunk, FileType

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.summarizer = None
        self.nlp = None
        self.encoder = None
        self._initialize_models()
    
    def _initialize_models(self):
        """Initialize AI models for processing"""
        try:
            # Initialize summarization model
            self.summarizer = pipeline(
                "summarization",
                model="facebook/bart-large-cnn",
                device=-1  # Use CPU
            )
            
            # Initialize spaCy for entity extraction
            self.nlp = spacy.load("en_core_web_sm")
            
            # Initialize sentence transformer for embeddings
            self.encoder = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
            
            logger.info("Document processor models initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing models: {e}")
            # Use fallback/dummy implementations
            self.summarizer = None
            self.nlp = None
            self.encoder = None
    
    async def process_file(self, file_path: str) -> Document:
        """Process uploaded file and extract content"""
        try:
            file_path_obj = Path(file_path)
            file_extension = file_path_obj.suffix.lower()
            
            # Extract text content
            content = await self._extract_text(file_path, file_extension)
            
            # Generate document metadata
            document_id = str(uuid.uuid4())
            title = self._extract_title(content, file_path_obj.stem)
            
            # Create document chunks
            chunks = self._create_chunks(content, document_id)
            
            # Generate summary
            summary = await self._generate_summary(content)
            
            # Extract tags/entities
            tags = self._extract_tags(content)
            
            # Get file stats
            file_stats = file_path_obj.stat()
            
            document = Document(
                id=document_id,
                title=title,
                content=content,
                summary=summary,
                tags=tags,
                file_type=self._get_file_type(file_extension),
                file_path=file_path,
                size_bytes=file_stats.st_size,
                uploaded_at=datetime.utcnow(),
                processed_at=datetime.utcnow(),
                chunk_ids=[chunk.id for chunk in chunks]
            )
            
            # Store chunks separately (would be handled by vector store)
            self.chunks = chunks  # Temporary storage
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise
    
    async def _extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_extension == '.txt' or file_extension == '.md':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            
            elif file_extension == '.pdf':
                return await self._extract_pdf_text(file_path)
            
            elif file_extension in ['.doc', '.docx']:
                return await self._extract_docx_text(file_path)
            
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            text = "Error: Could not extract text from PDF"
        
        return text.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return "Error: Could not extract text from DOCX"
    
    def _extract_title(self, content: str, filename: str) -> str:
        """Extract or generate document title"""
        lines = content.strip().split('\n')
        
        # Try to find a title in the first few lines
        for line in lines[:5]:
            line = line.strip()
            if line and len(line) < 100:  # Reasonable title length
                # Check if it looks like a title (no sentence-ending punctuation)
                if not line.endswith(('.', '!', '?')):
                    return line
        
        # Fallback to filename
        return filename.replace('_', ' ').replace('-', ' ').title()
    
    def _create_chunks(self, content: str, document_id: str, chunk_size: int = 500) -> List[DocumentChunk]:
        """Split document into chunks for vector storage"""
        words = content.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk_content = ' '.join(chunk_words)
            
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                document_id=document_id,
                content=chunk_content,
                chunk_index=i // chunk_size,
                word_count=len(chunk_words)
            )
            
            # Generate embedding if encoder is available
            if self.encoder:
                try:
                    embedding = self.encoder.encode(chunk_content).tolist()
                    chunk.embedding = embedding
                except Exception as e:
                    logger.warning(f"Failed to generate embedding: {e}")
            
            chunks.append(chunk)
        
        return chunks
    
    async def _generate_summary(self, content: str, max_length: int = 150) -> str:
        """Generate AI summary of document"""
        if not self.summarizer or len(content.split()) < 50:
            # Fallback: return first few sentences
            sentences = content.split('.')[:3]
            return '. '.join(sentences).strip() + '.'
        
        try:
            # Truncate content if too long for model
            max_input_length = 1000  # BART model limit
            truncated_content = ' '.join(content.split()[:max_input_length])
            
            summary = self.summarizer(
                truncated_content,
                max_length=max_length,
                min_length=50,
                do_sample=False
            )[0]['summary_text']
            
            return summary
        except Exception as e:
            logger.warning(f"Failed to generate AI summary: {e}")
            # Fallback summary
            return content[:200] + "..." if len(content) > 200 else content
    
    def _extract_tags(self, content: str) -> List[str]:
        """Extract tags/entities from content"""
        tags = []
        
        if self.nlp:
            try:
                # Process with spaCy
                doc = self.nlp(content[:5000])  # Limit length for performance
                
                # Extract named entities
                entities = [ent.text.lower() for ent in doc.ents 
                           if ent.label_ in ['PERSON', 'ORG', 'GPE', 'EVENT', 'WORK_OF_ART']]
                
                # Extract noun phrases as potential topics
                noun_phrases = [chunk.text.lower() for chunk in doc.noun_chunks 
                               if len(chunk.text.split()) <= 3]
                
                # Combine and filter
                all_candidates = list(set(entities + noun_phrases))
                tags = [tag for tag in all_candidates 
                       if len(tag) > 2 and len(tag) < 30][:10]  # Limit to 10 tags
                
            except Exception as e:
                logger.warning(f"Failed to extract tags with spaCy: {e}")
        
        # Fallback: simple keyword extraction
        if not tags:
            words = content.lower().split()
            # Simple frequency-based tag extraction
            word_freq = {}
            for word in words:
                if len(word) > 4 and word.isalpha():
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # Get most frequent words as tags
            tags = [word for word, freq in sorted(word_freq.items(), 
                                                key=lambda x: x[1], reverse=True)[:5]]
        
        return tags
    
    def _get_file_type(self, extension: str) -> FileType:
        """Convert file extension to FileType enum"""
        mapping = {
            '.pdf': FileType.PDF,
            '.txt': FileType.TXT,
            '.doc': FileType.DOC,
            '.docx': FileType.DOCX,
            '.md': FileType.MD
        }
        return mapping.get(extension.lower(), FileType.TXT)
    
    def health_check(self) -> Dict[str, Any]:
        """Check if document processor is healthy"""
        return {
            "status": "healthy" if self.encoder else "degraded",
            "models": {
                "summarizer": bool(self.summarizer),
                "nlp": bool(self.nlp),
                "encoder": bool(self.encoder)
            }
        }